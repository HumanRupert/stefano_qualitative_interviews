import typing as T
import uuid

import pandas as pd
import docx as dx


class Region(T.TypedDict):
    files_method: T.Callable[[], T.List[str]]
    name: str


QUESTIONS = [
    "What departments are responsible for recruitment at the MOE?",
    "Who recruits short term and contracted staff? What role does the MOE HR Directorate and its affiliate offices in provinces play in this recruitment process?",
    "What is the procedure for recruitment at the MOE? Who approves recruitments at provincial, district, and school level?",
    "What form of documentation is required to approve a recruitment at the MOE? Are any exceptions ever allowed? If yes, under what conditions?",
    "What is process of adding details of a newly recruited employee in the MOE payroll? Who gives final approval to this process?",
    "What documentation does this process require? Are any exceptions ever allowed? If yes, under what conditions?",
    "What is the process for requesting and approving transfers of teachers between schools?",
    "We are conducting a survey to better understand how to improve the MOE payroll. One aspect of this is to better understand the issue of “ghost workers.” Are you familiar with the term?",
    "Can you explain where ghost workers come from and what happens to the salary that is allocated to them?",
    "For every 100 employees on the MOE payroll, how many would you guess are “ghost” workers?",
    "What is the difference between a verified employee and ghost worker?",
    "Are there different types of ghost workers? How are they hired?",
    "Can a verified employee become a ghost worker at some point? If yes, under what conditions?",
    "Can a ghost worker become a verified employee at some point? If yes, how exactly?",
    "What is the benefit in being or in recruiting a ghost worker, especially if it involves legal action against them?",
    "Are ghost workers stand-alone individuals or they do have a source of support? If the later, where is this source of support located, in the government or outside the government?",
    "Is the impact of ghost workers on educational productivity significant enough to take legal action on it or should it be ignored?",
    "Does the existence of ghost workers affect educational service delivery? If yes, how exactly?",
    "If you think that educational service delivery would be improved by removing ghost workers, then who would oppose such a change and why?",
    "Can removing ghost workers result in increased school insecurity? If yes, how exactly?",
    "Will MOE officials in Kabul, at provincial education directorates, district departments or schools lose something if ghost workers are removed from the system?",
    "Will Mostofiats lose something if ghost workers are removed from the system?",
    "Will local power holders lose something if ghost workers are removed from the system?",
    "Do legitimate teachers benefit from the salaries paid by ghost workers?",
    "Do you think the MoE directs ghost worker salaries toward legitimate employees in order to encourage them to work in more challenging environments?",
    "In your assessment, how do you think that removing ghost workers might improve the situation in Afghanistan?",
    "In your assessment, how do you think that removing ghost workers might harm the situation in Afghanistan?",
]

CONSTRUAL_QUESTIONS = [
    "Did you introduce yourself, your affiliate organization to the interviewee",
    "Did the interviewee appear doubtful about your organizational affiliation and its linkages with government if any?",
    "Did the interviewee ask about the purpose of the interview? If yes, how did they receive your response to their question? Please also write your observations",
    "Was the interviewee ever reluctant in responding to questions? If yes, did they share why they were reluctant? If no, what was your observation of their reasons for reluctance?",
    "Did you inform the interviewee that responses will not be attributed to them personally? If yes, how they react to your answer?"
]


COLUMNS = ["Interview ID", "Month of Interview", "Numeric day of Month", "Year", "Province", "District", "Village", "Interviewer Name/ID", "Interviewer's Gender (M/F)", "Interviewer's age", "Interviewer's Current Profession", "Interviewer's Place of Birth", "Interviewer's Current Residing Place", "Interviewer's Ethnicity", "Interviewer's Highest Education  level",
           "Participant's name", "Gender (M/F)", "Age", "Education", "Marital Status(S/M)", "Occupation Position/Title", "Organization Affiliation", "Interviewee Province", "Interviewee District", "Interviewee Village/Nahia", "Ethnicity", "Tribe/Sub tribal Affiliation (Pashtuns Only)", "Religion Shia/Sunni", "Monthly Income", "Family Size"]


def _gen_df():
    columns = COLUMNS
    columns.extend([f"q{num+1}" for num in range(27)])
    columns.extend([f"cq{num+1}" for num in range(5)])
    df = pd.DataFrame(columns=columns)
    return df


def _add_table_cnt(row: T.List, doc: dx.Document):
    table = doc.tables[0]
    row.extend([cell.text for cell in table.rows[2].cells[2:5]])
    row.extend([cell.text for cell in table.rows[2].cells[6:]])
    row.extend(
        [cell.text for cell in table.columns[2].cells[3:11]])
    row.extend([cell.text for cell in table.columns[-1].cells[3:-3]])
    row.extend([cell.text for cell in table.columns[-1].cells[-2:]])


def _get_q_indices(paragraphs: T.List[str], qs: T.List[str]) -> T.List[int]:
    q_indices = []
    for q_ in qs:
        try:
            ix = [i for (i, v) in enumerate(paragraphs) if q_ in v][0]
        except:
            breakpoint()
        q_indices.append(ix)
    return q_indices


def _fix_inconsistent_qs(paragraphs):
    inconsistent_q = "Was the interviewee ever reluctant in responding to questions? If yes, did they share why"
    if(inconsistent_q in paragraphs):
        iq_ix = paragraphs.index(inconsistent_q)
        new_q = paragraphs[iq_ix] + " " + paragraphs[iq_ix+1]
        paragraphs[iq_ix] = new_q
        del paragraphs[iq_ix+1]


def _add_qs(row: T.List, doc: dx.Document, qs: T.List[str]):
    paragraphs = [
        p.text.replace("relunctant", "reluctant") for p in doc.paragraphs if not p.text.startswith("On")]

    _fix_inconsistent_qs(paragraphs)

    q_indices = _get_q_indices(paragraphs, qs)

    for ix, q_ix in enumerate(q_indices):
        try:
            ans_end_ix = paragraphs[q_ix:].index("") + q_ix \
                if ix == len(q_indices) - 1 \
                else q_indices[ix+1]
        except ValueError:
            ans_end_ix = len(paragraphs)
        ans = "".join(paragraphs[q_ix+1:ans_end_ix])
        row.append(ans)


def transform_load(region: Region):
    files = region["files_method"]()
    data = _gen_df()

    for file in files:
        doc = dx.Document(file)
        row = [str(uuid.uuid4())[:8]]
        _add_table_cnt(row, doc)
        _add_qs(row, doc, QUESTIONS)
        try:
            _add_qs(row, doc, CONSTRUAL_QUESTIONS)
        except:
            breakpoint()
        data.loc[len(data)] = row

    data.set_index("Interview ID", inplace=True)
    data.to_csv(f"out/{region['name']}.csv")
